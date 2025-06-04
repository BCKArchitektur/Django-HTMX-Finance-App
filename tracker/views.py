from django.shortcuts import render
from django.http import HttpResponse
from .models import Project , Logs , Contract , Section , Item , Task , ProjectPreset , UserPreset , Employee ,User , Client , SectionLibrary , Invoice
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from .forms import LogForm  , Hiddenform ,ProjectForm, ClientForm , ContractForm , InvoiceForm
from django.shortcuts import redirect
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, FloatField, Value
from django.db.models.functions import Cast
import datetime
from django.http import JsonResponse
from .models import Contract, Section , Task, Item
from django.utils import timezone
import pytz
from django.db.models import Sum, FloatField, F
from django.db.models import DateField
import json
from .models import ProjectPreset
from django.http import JsonResponse, HttpResponseBadRequest
from django.utils.decorators import method_decorator
from django.views import View
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from .forms import AddUsersForm, AddBudgetForm
from django.contrib import messages
from django.urls import reverse
from urllib.parse import urlencode
import docx
from docxtpl import DocxTemplate
import os
from django.conf import settings
from datetime import date
from django.urls import reverse
from django.http import HttpResponseRedirect
from django.http import JsonResponse, Http404  # Import Http404
from decimal import Decimal  # Import Decimal for precision handling
from datetime import datetime
from django.db.models import JSONField

@login_required
def toggle_dark_mode(request):
    user = request.user
    user.dark_mode = not user.dark_mode
    user.save()
    return JsonResponse({'dark_mode': user.dark_mode})

    
def login_page(request):
    return render(request,'projects/login_page.html')

def logout_page(request):
    return render(request,'projects/logout_page.html')


@login_required
def load_contracts(request):
    project_id = request.GET.get('project_id')
    if not project_id or not project_id.isdigit():
        return JsonResponse({'contracts': []})
    
    contracts = Contract.objects.filter(project__id=int(project_id), user=request.user).order_by('contract_name')
    contract_list = [{'id': contract.id, 'contract_name': contract.contract_name} for contract in contracts]
    return JsonResponse({'contracts': contract_list})


@login_required
def load_sections(request):
    contract_id = request.GET.get('contract_id')
    if not contract_id or not contract_id.isdigit():
        return JsonResponse({'sections': []})
    
    sections = Section.objects.filter(contract__id=int(contract_id), user=request.user).order_by('section_name')
    section_list = [{'id': section.id, 'section_name': section.section_name} for section in sections]
    return JsonResponse({'sections': section_list})


@login_required
def load_Items(request):
    section_id = request.GET.get('section_id')
    if not section_id or not section_id.isdigit():
        return JsonResponse({'Items': []})
    
    items = Item.objects.filter(section__id=int(section_id), users=request.user)
    item_list = [{'id': item.id, 'Item_name': item.Item_name} for item in items]
    return JsonResponse({'Items': item_list})

def load_tasks(request):
    item_id = request.GET.get('Item_id')
    
    # Check if Item_id is valid
    if not item_id or not item_id.isdigit():
        return HttpResponseBadRequest("Invalid Item ID")
    
    item_id = int(item_id)
    
    try:
        # Retrieve the Item object
        item = Item.objects.get(id=item_id)
        # Get the related tasks through the many-to-many relationship
        tasks = item.tasks.all()
        
        # Prepare the task list
        task_list = [{'id': task.id, 'task_name': task.task_name} for task in tasks]
        
        return JsonResponse({'tasks': task_list})
    
    except Item.DoesNotExist:
        return HttpResponseBadRequest("Item not found")


@login_required
def delete_log(request, log_id):
    
    # Force clean up of the log_id to remove commas
    try:
        cleaned_log_id = ''.join(log_id.split(','))  # Remove commas by splitting and rejoining
        cleaned_log_id = int(cleaned_log_id)  # Convert to integer
    except ValueError as e:
        return JsonResponse({'success': False, 'error': 'Invalid log_id format'}, status=400)
    
    if request.method == 'POST':
        try:
            log = get_object_or_404(Logs, id=cleaned_log_id, user=request.user)  # Ensure the log belongs to the user
            
            log.delete()
            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    else:
        return JsonResponse({'success': False}, status=400)


@login_required
def index(request):
    context = {
    }
    return render(request, 'tracker/index.html', context)

@login_required
def projects(request):
    projects = Project.objects.all()  # Show all projects
    clients = Client.objects.all()  # Assuming all clients should be shown
    project_form = ProjectForm()
    client_form = ClientForm()

    if request.method == 'POST':
        if 'project_name' in request.POST:
            project_form = ProjectForm(request.POST)
            if project_form.is_valid():
                project_form.save()
                return redirect('projects')
        elif 'client_name' in request.POST:
            client_form = ClientForm(request.POST)
            if client_form.is_valid():
                client_form.save()
                return redirect('projects')

    context = {
        'projects': projects.order_by('-project_no'),
        'clients': clients.order_by('-firm_name'),
        'project_form': project_form,
        'client_form': client_form,
    }
    return render(request, 'tracker/projects.html', context)


@csrf_exempt
@login_required
def delete_project(request, project_id):
    if request.method == 'POST':
        project = get_object_or_404(Project, id=project_id)
        project.delete()
        return JsonResponse({'status': 'success'})
    return HttpResponseBadRequest("Invalid request")

@csrf_exempt
@login_required
def delete_client(request, client_id):
    if request.method == 'POST':
        client = get_object_or_404(Client, id=client_id)
        client.delete()
        return JsonResponse({'status': 'success'})
    return HttpResponseBadRequest("Invalid request")

@login_required
def edit_project(request, project_id):
    project = get_object_or_404(Project, id=project_id)

    # Store the old project name for later use (if the project name is updated)
    old_project_name = project.project_name

    # Fetch all projects to pass to the context
    all_projects = Project.objects.all()

    contracts = project.contract.all()
    clients = Client.objects.all()
    users = User.objects.all()
    contract_form = ContractForm()
    section_library = SectionLibrary.objects.all()  # Fetch all sections from the library
    invoices = Invoice.objects.filter(project=project)

    if request.method == 'POST':
        if 'project_name' in request.POST:
            # Handle project name update
            new_project_name = request.POST.get('project_name')
            if old_project_name != new_project_name:  # Only proceed if the name has actually changed
                # Update the project name
                project.name = new_project_name
                project.save()

                # Find all logs with the old project name and update them
                Logs.objects.filter(log_project_name=old_project_name).update(log_project_name=new_project_name)

            return handle_project_form(request, project)
        elif 'contract_name' in request.POST and not request.POST.get('contract_id'):
            return handle_new_contract_form(request, project)
        elif request.POST.get('contract_id'):
            return handle_existing_contract_form(request, project)
        elif 'user' in request.POST:
            updated_users = request.POST.getlist('user')
            handle_user_updates(project, updated_users)
    else:
        form = ProjectForm(instance=project)

    # Get the VAT percentages for each contract
    vat_percentages = {contract.id: contract.vat_percentage for contract in contracts}

    context = {
        'form': form,
        'project': project,
        'contracts': contracts,
        'contract_form': contract_form,
        'clients': clients,
        'users': users,
        'section_library': section_library, 
        'invoices': invoices.order_by('-created_at'),
        'vat_percentages': vat_percentages,  # Pass VAT percentages to the template
        'all_projects': all_projects
    }
    return render(request, 'tracker/edit_project.html', context)




def handle_user_updates(project, updated_users):
    current_users = set(project.user.all())
    updated_users_set = set(User.objects.filter(id__in=updated_users))

    # Users to be added and removed
    users_to_add = updated_users_set - current_users
    users_to_remove = current_users - updated_users_set

    # Add users to project
    for user in users_to_add:
        project.user.add(user)
        for contract in project.contract.all():
            contract.user.add(user)
            for section in contract.section.all():
                section.user.add(user)
                for item in section.Item.all():
                    item.users.add(user)

    # Remove users from project
    for user in users_to_remove:
        project.user.remove(user)
        for contract in project.contract.all():
            contract.user.remove(user)
            for section in contract.section.all():
                section.user.remove(user)
                for item in section.Item.all():
                    item.users.remove(user)
    
    project.save()


def get_library_section(request, section_id=None):
    """
    Fetches either:
    - A specific section by ID (if section_id is provided)
    - Multiple sections by LP if 'lp' parameter is given
    """
    
    lp_param = request.GET.get("lp")  # Get LP filter if provided

    if lp_param:  # If LP is passed in the request, return multiple sections
        sections = SectionLibrary.objects.filter(lp_identifier=lp_param)
        sections_data = [
            {
                'section_name': section.name,
                'items': [
                    {
                        'Item_name': item.name,
                        'tasks': [{'task_name': task.name} for task in item.tasks.all()]
                    }
                    for item in section.items.all()
                ]
            }
            for section in sections
        ]
        return JsonResponse(sections_data, safe=False)
    
    # If no LP is provided, return a specific section by ID (original behavior)
    if section_id:
        section = get_object_or_404(SectionLibrary, id=section_id)
        section_data = {
            'section_name': section.name,
            'items': [
                {
                    'Item_name': item.name,
                    'tasks': [{'task_name': task.name} for task in item.tasks.all()]
                }
                for item in section.items.all()
            ]
        }
        return JsonResponse(section_data)

    return JsonResponse({"error": "Invalid request, provide section_id or lp parameter"}, status=400)


def handle_project_form(request, project):
    form = ProjectForm(request.POST, instance=project)
    if form.is_valid():
        form.save()
        messages.success(request, "Project updated successfully.")
    else:
        messages.error(request, "Error updating project: {}".format(form.errors))
    return redirect('edit_project', project_id=project.id)


import json
from django.shortcuts import redirect
from django.contrib import messages
from .models import Contract, Section, Item, Task


def handle_existing_contract_form(request, project):
    print("Request received for handling existing contract form.")

    # Retrieve contract ID
    contract_id = request.POST.get('contract_id')
    if not contract_id:
        print("No contract ID provided.")
        messages.error(request, "Contract ID is missing.")
        return redirect('edit_project', project_id=project.id)

    try:
        contract = Contract.objects.get(id=contract_id)
        print(f"Contract found: {contract}")
    except Contract.DoesNotExist:
        print(f"Contract with ID {contract_id} does not exist.")
        messages.error(request, "Contract not found.")
        return redirect('edit_project', project_id=project.id)
    
    # Retrieve and parse HOAI data
    hoai_data = request.POST.get('hoai_data', '{}').strip()  # Remove unnecessary whitespace
    if not hoai_data:
        hoai_data = '{}'
    try:
        hoai_data_parsed = json.loads(hoai_data)
    except json.JSONDecodeError:
        hoai_data_parsed = {}  # Set default empty dictionary
        print("Error decoding HOAI data JSON")  # Debugging

    try:
        hoai_data_parsed = json.loads(hoai_data)
    except json.JSONDecodeError:
        hoai_data_parsed = {}
        print("Error decoding HOAI data JSON")  # Debugging line

    # Validate contract form
    contract_form = ContractForm(request.POST, instance=contract)
    if not contract_form.is_valid():
        print(f"Contract form errors: {contract_form.errors}")
        messages.error(request, f"Error updating contract: {contract_form.errors}")
        return redirect('edit_project', project_id=project.id)

    # Save the contract form
    contract = contract_form.save()
    print("Contract form saved.")

    # Parse contract JSON
    contract_json = request.POST.get('contract_json')
    print('contract_json:', contract_json)
    if not contract_json:
        print("No contract JSON data provided.")
        messages.error(request, "No contract JSON data provided.")
        return redirect('edit_project', project_id=project.id)

    try:
        contract_data = json.loads(contract_json)
        print(f"Parsed contract data: {contract_data}")
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        messages.error(request, "Error decoding the contract JSON data.")
        return redirect('edit_project', project_id=project.id)

    # Process sections
    sections_to_keep = []
    for section_data in sorted(contract_data.get('sections', []), key=lambda x: x.get('order', 0)):  # Sort sections by order
        print(f"Processing section: {section_data}")
        section_id = section_data.get('id')  
        section_name = section_data['section_name']
        section_billed_hourly = section_data.get('section_billed_hourly', False)
        section_order = section_data.get('order', 0)  

        # Retrieve or create section
        if section_id and section_id.isdigit():  
            section = Section.objects.filter(id=section_id).first()
            if section:
                print(f"Updating existing section: {section}")
                section.section_name = section_name
                section.section_billed_hourly = section_billed_hourly
                section.order = section_order  
                section.save()
            else:
                print(f"Section ID {section_id} provided, but no section found! Creating new.")
                section = Section.objects.create(
                    section_name=section_name,
                    section_billed_hourly=section_billed_hourly,
                    order=section_order
                )
        else:
            print(f"-Creating a new section: {section_name}")
            section = Section.objects.create(
                section_name=section_name,
                section_billed_hourly=section_billed_hourly,
                order=section_order
            )
            section_id = section.id

        print(f"Processed section: {section}")
        sections_to_keep.append(section)

        # Process items in the section
        items_to_keep = []
        for item_data in sorted(section_data.get('items', []), key=lambda x: x.get('order', 0)):  
            print(f"Processing item: {item_data}")
            item_id = item_data.get('id')
            item_name = item_data['Item_name']
            description = item_data.get('description', '')
            item_order = item_data.get('order', 0)  

            # Retrieve or create item
            if item_id and not str(item_id).startswith('new-'):
                item = Item.objects.filter(id=item_id).first()
                if item:
                    print(f"Updating existing item: {item}")
                    item.Item_name = item_name
                    item.description = description
                    item.quantity = item_data.get('quantity', item.quantity)
                    item.unit = item_data.get('unit', item.unit)
                    item.rate = item_data.get('rate', item.rate)
                    item.order = item_order  
                    item.save()
                else:
                    print(f"Item with ID {item_id} not found. Creating new item.")
                    item = Item.objects.create(
                        Item_name=item_name,
                        description=description,
                        quantity=item_data.get('quantity', 0),
                        unit=item_data.get('unit', 'Std'),
                        rate=item_data.get('rate', 0.0),
                        order=item_order
                    )
            else:
                print(f"Creating a new item: {item_name}")
                item = Item.objects.create(
                    Item_name=item_name,
                    description=description,
                    quantity=item_data.get('quantity', 0),
                    unit=item_data.get('unit', 'Std'),
                    rate=item_data.get('rate', 0.0),
                    order=item_order
                )

            print(f"Processed item: {item}")
            items_to_keep.append(item)

            # Process tasks for the item
            tasks_to_keep = []
            for task_data in item_data.get('tasks', []):
                print(f"Processing task: {task_data}")
                task_name = task_data['task_name']

                # Create task (no ID means tasks are always new)
                task = Task.objects.create(task_name=task_name)
                print(f"Created task: {task}")
                tasks_to_keep.append(task)

            item.tasks.set(tasks_to_keep)  
            section.Item.add(item)  

        print(f"Setting items for section: {section.section_name}")
        section.Item.set(items_to_keep)  

    print("Setting sections for the contract.")
    contract.section.set(sections_to_keep)  
    contract.save()

    # Save contract with updated HOAI data
    if contract.hoai_data:
        contract.hoai_data = hoai_data_parsed
        contract.save()

    if contract.hoai_data:
        hoai_data = json.loads(request.POST.get('hoai_data', '{}'))
        zuschlag_value = hoai_data.get("zuschlag", 0)

        contract.zuschlag_value = zuschlag_value
        contract.vat_percentage=Decimal(hoai_data.get("vat", 19.00))

        contract.additional_fee_percentage =float( hoai_data.get("nebenKosten", 6.5))
        contract.save()

    #  Assign Budget if HOAI Mode is Enabled
    if hoai_data_parsed:
        assign_budget_to_contract(contract, hoai_data_parsed)

    # Associate contract with project
    project.contract.add(contract)
    project.save()

    print("Contract and project updated successfully.")
    messages.success(request, "Contract updated successfully.")

    return redirect('edit_project', project_id=project.id)

from decimal import Decimal
import re


import re
from decimal import Decimal

def parse_german_number(number_string):
    """
    Converts a German-formatted number (e.g., "101.244,50") into a Decimal (101244.50).
    """
    try:
        number_string = str(number_string)
        normalized_number = number_string.replace('.', '').replace(',', '.')
        return Decimal(normalized_number)  
    except (ValueError, AttributeError):
        return Decimal(0)



def assign_budget_to_contract(contract, hoai_data):
    """
    Assigns a budget to contract sections based on HOAI data.
    Looks for LP sections and assigns budget to all items within the section.
    If it's an LP section, it assumes only one item inside.
    """
    grundhonorar_raw = hoai_data.get("grundhonorar", "0")
    grundhonorar = parse_german_number(grundhonorar_raw) 

    print(f"🔹 Assigning Budget - Grundhonorar: {grundhonorar}")



    for section in contract.section.all():
        section_name = section.section_name.strip()

        # Extract LP key from section name
        lp_match = re.search(r"LP(\d+)", section_name, re.IGNORECASE)

        if lp_match:
            lp_key = f"lp{lp_match.group(1)}"  # Convert to format "lp1", "lp2"
            
            if lp_key in hoai_data.get("lp_values", {}):
                lp_percentage = Decimal(hoai_data["lp_values"][lp_key]) 
                
                # Process all items in the section
                items = list(section.Item.all())
                
                if items:  # Ensure there are items
                    if len(items) == 1:
                        item = items[0]  
                    else:
                        print(f"!! Multiple items found in LP section !!'{section_name}', processing all.")

                    for item in items:
                        # Assign Budget Values
                        item.quantity = lp_percentage
                        item.unit = "%"  
                        item.rate = grundhonorar
                        item.total = (lp_percentage / Decimal(100)) * grundhonorar  
                        item.save()

                        print(f"Budget Assigned: {item.Item_name} | Qty: {lp_percentage}% | Rate: {grundhonorar} | Total: {item.total}")

    print("-Budget assignment completed.")



from django.http import JsonResponse
from django.shortcuts import redirect
import json

def handle_new_contract_form(request, project):
    contract_name = request.POST.get('contract_name')
    contract_json = request.POST.get('contract_json')
    contract_no = request.POST.get('contract_no')

    # Retrieve all users associated with the project
    user_ids = project.user.values_list('id', flat=True)

    print("POST data:", request.POST)  # Debugging

    # Retrieve and clean hoai_data
    hoai_data_str = request.POST.get('hoai_data', '').strip()

    # Initialize hoai_data as an empty dictionary
    hoai_data = {}

    # Only parse if hoai_data_str is not empty
    if hoai_data_str:
        try:
            hoai_data = json.loads(hoai_data_str)
        except json.JSONDecodeError:
            print("Error decoding HOAI data JSON")  # Debugging


  
    zuschlag_value = hoai_data.get("zuschlag", 0)

    # Create Contract object
    contract = Contract.objects.create(
        contract_name=contract_name,
        contract_no=contract_no,
        hoai_data=hoai_data, 
        zuschlag_value=zuschlag_value,
        vat_percentage=hoai_data.get("vat"),
        additional_fee_percentage = hoai_data.get("nebenKosten"),
    )
    print('creating new contract with nebenkosten',hoai_data.get("nebenKosten", 6.5))
    contract.user.set(user_ids)

    print("Received contract JSON:", contract_json)  # Debugging

    if contract_json:
        try:
            contract_data = json.loads(contract_json)
            print("Parsed contract data:", contract_data)  # Debugging

            for section_data in contract_data.get('sections', []):
                section_name = section_data['section_name']
                section_billed_hourly = section_data.get('section_billed_hourly', False)

                section = Section.objects.create(
                    section_name=section_name,
                    section_billed_hourly=section_billed_hourly
                )
                print("Processed section:", section_name)  # Debugging

                section.user.set(user_ids)

                for item_data in section_data.get('items', []):
                    Item_name = item_data['Item_name']
                    description = item_data.get('description', '')

                    item = Item(
                        Item_name=Item_name,
                        description=description
                    )
                    item.set_project_context(project)  
                    item.save()
                    item.users.set(user_ids)


                    for task_data in item_data.get('tasks', []):
                        task_name = task_data['task_name']
                        task = Task.objects.create(task_name=task_name)
                        print("Processed task:", task_name)  # Debugging
                        item.tasks.add(task)

                    section.Item.add(item)

                contract.section.add(section)

            contract.save()
            project.contract.add(contract)
            project.save()

            messages.success(request, "New contract added successfully.")

            # Assign Budget if HOAI Mode is Enabled
            if hoai_data:
                assign_budget_to_contract(contract, hoai_data)

        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}")  # Debugging
            messages.error(request, "Error decoding the contract JSON data.")

    else:
        messages.error(request, "No contract JSON data provided.")

    # Return JSON if AJAX request
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({"contract_id": contract.id})

    # Otherwise, redirect for normal form submission
    return redirect('edit_project', project_id=project.id)



def update_contract_details(request, contract):
    section_names = request.POST.getlist('section_name')
    Item_names = request.POST.getlist('Item_name')
    task_names = request.POST.getlist('task_name')

    contract.section.clear()
    for section_name in section_names:
        section, _ = Section.objects.get_or_create(section_name=section_name)
        section.Item.clear()
        for Item_name in Item_names:
            item, _ = Item.objects.get_or_create(Item_name=Item_name)
            item.tasks.clear()
            for task_name in task_names:
                task, _ = Task.objects.get_or_create(task_name=task_name)
                item.tasks.add(task)
            section.Item.add(item)
        contract.section.add(section)
    contract.save()


@login_required
def edit_client(request, client_id):
    client = get_object_or_404(Client, id=client_id)
    if request.method == 'POST':
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            # Construct the URL with the query parameter
            url = reverse('projects') + '?' + urlencode({'tab': 'clients'})
            return redirect(url)
    else:
        form = ClientForm(instance=client)
    return render(request, 'tracker/edit_client.html', {'form': form, 'client': client})



@login_required
def dashboard(request):
    context = {
    }
    return render(request, 'tracker/dashboard.html', context)

#view to get current time
def get_berlin_time():
    berlin_timezone = pytz.timezone('Europe/Berlin')
    current_time = timezone.now()
    localized_time = current_time.astimezone(berlin_timezone)
    formatted_time = localized_time.strftime('%d/%m/%y %H:%M:%S')
    return formatted_time

#Log view to create log
@login_required
def log_create_compact(request):
    projects = Project.objects.filter(user=request.user)
    project_presets = ProjectPreset.objects.filter(user=request.user)
    logs = Logs.objects.filter(user=request.user).order_by('-log_timestamps')[:100]

    today = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).strftime('%Y-%m-%d')
    logs_today = Logs.objects.filter(user=request.user, log_timestamps__startswith=today)
    total_hours_today = logs_today.annotate(
        numeric_time=Cast('log_time', FloatField())
    ).aggregate(total_time=Sum('numeric_time'))['total_time'] or 0

    # Get the employee's assigned hours
    employee = get_object_or_404(Employee, user=request.user)
    
    # Determine today's day of the week
    day_of_week = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).weekday()

    # Mapping of day of the week to the corresponding hours assigned variable
    hours_assigned_mapping = {
        0: employee.hours_assigned_monday,
        1: employee.hours_assigned_tuesday,
        2: employee.hours_assigned_wednesday,
        3: employee.hours_assigned_thursday,
        4: employee.hours_assigned_friday,
    }

    hours_assigned_today = hours_assigned_mapping.get(day_of_week, 0)

    # Calculate the progress percentage
    progress_percentage = (total_hours_today / hours_assigned_today) * 100 if hours_assigned_today > 0 else 0

    form = Hiddenform(request.POST or None, user=request.user)
    if request.method == 'POST' and form.is_valid():
        log_project_name = form.cleaned_data['log_project_name']
        log_contract_id = form.cleaned_data['log_contract'].replace(",", "") # Use the contract ID
        log_tasks = form.cleaned_data['log_tasks']
        log_section_id = form.cleaned_data['log_section'].replace(",", "") # Use the section ID
        print(f"log_section_id: {log_section_id}")  # Debugging line
        log_item_id = form.cleaned_data['log_Item'].replace(",", "")  # Use the item ID
        log_time = form.cleaned_data['log_time']
        log_timestamps = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).strftime('%Y-%m-%d %H:%M:%S')
        
        # Fetch by ID
        log_project = get_object_or_404(Project, project_name=log_project_name)
        log_contract = get_object_or_404(Contract, id=log_contract_id)  # Fetch by ID instead of name
        log_section = get_object_or_404(Section, id=log_section_id)  # Fetch by ID instead of name
        log_item = get_object_or_404(Item, id=log_item_id)  # Fetch by ID instead of name

        log_entry = Logs.objects.create(
            log_project_name=log_project_name,
            log_contract=log_contract,
            log_section=log_section,
            log_Item=log_item,
            log_time=log_time,
            log_timestamps=log_timestamps,
            log_tasks=log_tasks,
            user=request.user
        )
        log_entry.save()
        return redirect('log_create_compact')

    context = {
        'project_presets': project_presets,
        'form': form,
        'logs': logs,
        'total_hours_today': total_hours_today,
        'hours_assigned_today': hours_assigned_today,
        'progress_percentage': progress_percentage,
        'projects': projects,
        'employee': employee,
    }
    return render(request, 'tracker/log_create_compact.html', context)



@login_required
def log_create(request):
    projects = Project.objects.filter(user=request.user, status='0')
    logs = Logs.objects.filter(user=request.user).order_by('-log_timestamps')[:100]
    today = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).strftime('%Y-%m-%d')
    logs_today = Logs.objects.filter(user=request.user, log_timestamps__startswith=today)
    total_hours_today = logs_today.annotate(
        numeric_time=Cast('log_time', FloatField())
    ).aggregate(total_time=Sum('numeric_time'))['total_time'] or 0

    # Get the employee's assigned hours
    employee = get_object_or_404(Employee, user=request.user)
    
    # Determine today's day of the week
    day_of_week = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).weekday()

    # Mapping of day of the week to the corresponding hours assigned variable
    hours_assigned_mapping = {
        0: employee.hours_assigned_monday,
        1: employee.hours_assigned_tuesday,
        2: employee.hours_assigned_wednesday,
        3: employee.hours_assigned_thursday,
        4: employee.hours_assigned_friday,
    }

    hours_assigned_today = hours_assigned_mapping.get(day_of_week, 0)

    # Calculate the progress percentage
    progress_percentage = (total_hours_today / hours_assigned_today) * 100 if hours_assigned_today > 0 else 0

    form = LogForm(request.POST or None, user=request.user)
    if request.method == 'POST' and form.is_valid():
        log_project_name = form.cleaned_data['log_project_name'].project_name
        log_contract = form.cleaned_data['log_contract']
        log_tasks = form.cleaned_data['log_tasks']
        log_section = form.cleaned_data['log_section']
        print(f"log_section: {log_section}")  # Debugging line
        log_Item = form.cleaned_data['log_Item']
        log_time = form.cleaned_data['log_time']


        # Get the submitted date and keep the current time
        submitted_date = form.cleaned_data['log_date'] or today  # Fallback to today's date if empty
        current_time = timezone.now().astimezone(pytz.timezone('Europe/Berlin')).strftime('%H:%M:%S')
        log_timestamps = f"{submitted_date} {current_time}"
        
        # Create the log entry
        log_entry = Logs.objects.create(
            log_project_name=log_project_name,
            log_contract=log_contract,
            log_section=log_section,
            log_Item=log_Item,
            log_time=log_time,
            log_timestamps=log_timestamps,
            log_tasks=log_tasks,
            user=request.user
        )
        log_entry.save()

        project = form.cleaned_data['log_project_name']
        default_contract = form.cleaned_data['log_contract']
        default_section = form.cleaned_data['log_section']
        default_Item = form.cleaned_data['log_Item']

        user_presets = ProjectPreset.objects.filter(user=request.user)
        if user_presets.count() >= 4:
            # Delete the oldest preset specific to the user
            oldest_preset = user_presets.order_by('id').first()
            oldest_preset.delete()

        # Check if an existing preset with the same values already exists
        existing_preset = ProjectPreset.objects.filter(
            user=request.user,
            project=project,
            default_contract=default_contract,
            default_section=default_section,
            default_Item=default_Item
        ).exists()

        # Only create a new preset if no existing one is found
        if not existing_preset:
            ProjectPreset.objects.create(
                user=request.user,
                project=project,
                default_contract=default_contract,
                default_section=default_section,
                default_Item=default_Item,
            )

        return redirect('log_create')
    else:
        form = LogForm(user=request.user)

    context = {
        'form': form,
        'logs': logs,
        'total_hours_today': total_hours_today,
        'hours_assigned_today': hours_assigned_today,
        'progress_percentage': progress_percentage,
        'projects': projects.order_by('-project_no'),
        'date_override': employee.date_override,
    }
    return render(request, 'tracker/log_create.html', context)




@login_required
def project_details(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    contracts = project.contract_set.all()

    contract_data = []
    for contract in contracts:
        sections = contract.section_set.all()
        for section in sections:
            # Mock data for used budget, replace with actual logic
            used_budget = section.allocated_budget * 1.2  # Example logic
            contract_data.append({
                'contract_name': contract.contract_name,
                'section_name': section.section_name,
                'allocated_budget': section.allocated_budget,
                'used_budget': used_budget
            })

    response = {
        'project_name': project.project_name,
        'contracts': contract_data
    }

    if request.is_ajax():
        return JsonResponse(response)
    else:
        return render(request, 'tracker/project_details.html', {'project': project, 'contracts': contract_data})


@login_required
def add_project(request):
    if request.method == 'POST':
        project_name = request.POST.get('project_name')
        project_address = request.POST.get('project_address', '')
        new_project = Project.objects.create(
            user=request.user,
            project_name=project_name,
            project_address=project_address,
            project_no="PN-" + str(Project.objects.count() + 1)  # Example project number
        )
        return redirect('projects')
    return redirect('projects')


from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Contract, Invoice, User

def get_contract_scope(request, contract_id):
    contract = get_object_or_404(Contract, id=contract_id)
    return JsonResponse({
        'contract_id': contract.id,
        'scope': contract.scope_of_work if contract.scope_of_work else ''
    })

def load_contract_data(request):
    contract_id = request.GET.get('contract_id')
    contract = get_object_or_404(Contract, id=contract_id)

    # Ensure the contract belongs to at least one project
    project_id = contract.project_set.first().id if contract.project_set.exists() else None
    if not project_id:
        return JsonResponse({'error': 'Contract does not belong to any project'}, status=400)

    # Fetch all previous invoices for this contract
    previous_invoices = Invoice.objects.filter(project_id=project_id, contract_id=contract_id).order_by('-created_at')

    invoices_exist = previous_invoices.exists()
    latest_invoice = previous_invoices.first()  # Fetch the most recent invoice
    is_cumulative = latest_invoice.is_cumulative if latest_invoice else False

    # Retrieve the latest provided quantities if cumulative mode
    latest_provided_quantities = latest_invoice.provided_quantities if is_cumulative and latest_invoice else {}

    users = list(User.objects.all().values('id', 'username'))

    sections = sorted(contract.section.all(), key=lambda s: getattr(s, 'order', 0))

    section_data = []

    for section in sections:
        items = sorted(section.Item.all(), key=lambda i: getattr(i, 'order', 0))

        item_data = []
        for item in items:
            if is_cumulative:
                previous_provided_quantity = latest_provided_quantities.get(str(item.id), {}).get('quantity', 0) if latest_invoice else 0
            else:
                previous_provided_quantity = sum(
                    invoice.provided_quantities.get(str(item.id), {}).get('quantity', 0)
                    for invoice in previous_invoices
                )

            available_quantity = item.quantity - previous_provided_quantity

            # Calculate hours_logged from Logs model
            hours_logged = Logs.objects.filter(log_Item=item).aggregate(
                total_hours=Sum('log_time')
            )['total_hours'] or 0

            item_data.append({
                'order': getattr(item, 'order', 0),
                'id': item.id,
                'Item_name': item.Item_name,
                'description': item.description,
                'quantity': item.quantity,
                'available_quantity': available_quantity,
                'previous_provided_quantity': previous_provided_quantity,  
                'unit': item.unit,
                'rate': item.rate,
                'total': item.total,
                'users': list(item.users.values_list('id', flat=True)),
                'tasks': list(item.tasks.values('id', 'task_name')),
                'hours_logged': hours_logged,  
            })

        section_data.append({
            'order': getattr(section, 'order', 0),
            'id': section.id,
            'section_name': section.section_name,
            'section_billed_hourly': section.section_billed_hourly,
            'exclude_from_nachlass': section.exclude_from_nachlass,
            'items': item_data,
        })

    hoai_data = contract.hoai_data if contract.hoai_data else {}
    zuschlag_value = contract.zuschlag_value
    nachlass_value = contract.nachlass_value
    nachlass_percentage = contract.nachlass_percentage


    contract_data = {
        'contract_name': contract.contract_name,
        'contract_no': contract.contract_no,
        'users': users,
        'sections': section_data,
        'additional_fee_percentage': contract.additional_fee_percentage,
        'vat_percentage': contract.vat_percentage,
        'invoices_exist': invoices_exist,
        'is_cumulative': is_cumulative, 
        'latest_provided_quantities': latest_provided_quantities,
        'hoai_data': hoai_data,
        'zuschlag_value': zuschlag_value,
        'nachlass_value' : nachlass_value,
        'nachlass_percentage' : nachlass_percentage,

    }

    return JsonResponse(contract_data)


def check_task_name(request):
    task_name = request.GET.get('task_name', None)
    is_taken = Task.objects.filter(task_name__iexact=task_name).exists()
    data = {
        'is_taken': is_taken,
        'message': 'Task name is already taken.' if is_taken else 'Task name is available.'
    }
    return JsonResponse(data)

def check_section_name(request):
    section_name = request.GET.get('section_name', None)
    is_taken = Section.objects.filter(section_name__iexact=section_name).exists()
    data = {
        'is_taken': is_taken,
        'message': 'Section name is already taken.' if is_taken else 'Section name is available.'
    }
    return JsonResponse(data)

def check_Item_name(request):
    Item_name = request.GET.get('Item_name', None)
    is_taken = Item.objects.filter(Item_name__iexact=Item_name).exists()
    data = {
        'is_taken': is_taken,
        'message': 'Item name is already taken.' if is_taken else 'Item name is available.'
    }
    return JsonResponse(data)

def check_contract_name(request):
    contract_name = request.GET.get('contract_name', None)
    is_taken = Contract.objects.filter(contract_name__iexact=contract_name).exists()
    data = {
        'is_taken': is_taken,
        'message': 'Contract name is already taken.' if is_taken else 'Contract name is available.'
    }
    return JsonResponse(data)


@csrf_exempt
@login_required
def add_users(request):
    if request.method == 'POST':
        contract_id = request.POST.get('contract_id')
        contract = get_object_or_404(Contract, id=contract_id)
        
        for section in contract.section.all():
            for item in section.Item.all():
                for user in User.objects.all():
                    user_checkbox = f'user_item_{user.id}_{item.id}'
                    if user_checkbox in request.POST:
                        item.users.add(user)
                    else:
                        item.users.remove(user)

        # Now update the users of the sections based on the users of their items
        for section in contract.section.all():
            section_users = set()
            for item in section.Item.all():
                section_users.update(item.users.all())
            section.user.set(section_users)

        # Update the users of the contract based on the users of its sections
        contract_users = set()
        for section in contract.section.all():
            contract_users.update(section.user.all())
        contract.user.set(contract_users)

        return JsonResponse({'status': 'success'})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)


@csrf_exempt
@login_required
def add_budget(request):
    if request.method == 'POST':
        contract_id = request.POST.get('contract_id')
        if not contract_id:
            return JsonResponse({'error': 'Missing contract ID'}, status=400)

        contract = get_object_or_404(Contract, id=contract_id)

        # Handle additional fee percentage
        additional_fee_percentage = request.POST.get('additional_fee_percentage')
        try:
            contract.additional_fee_percentage = float(additional_fee_percentage)
        except (ValueError, TypeError):
            contract.additional_fee_percentage = 0.0

        # Handle VAT percentage
        vat_percentage = request.POST.get('vat_percentage')
        try:
            contract.vat_percentage = float(vat_percentage)
        except (ValueError, TypeError):
            contract.vat_percentage = 0.0

        # Handle Nachlass value
        nachlass_value = request.POST.get('nachlass_value')
        try:
            contract.nachlass_value = float(nachlass_value)
        except (ValueError, TypeError):
            contract.nachlass_value = 0.0
        # Handle Nachlass percentage
        nachlass_percentage = request.POST.get('nachlass_percentage')
        try:
            contract.nachlass_percentage = float(nachlass_percentage)
        except (ValueError, TypeError):
            contract.nachlass_percentage = 0.0

        # Handle the sections and items
        for section in contract.section.all():
            field_name = f'exclude_section_{section.id}'
            section.exclude_from_nachlass = bool(request.POST.get(field_name) == 'on')
            print(f"Section {section.id}- {section.section_name}  exclude_from_nachlass: {section.exclude_from_nachlass}") 
            section.save()
            for item in section.Item.all():
                quantity_key = f'quantity_{item.id}'
                unit_key = f'unit_{item.id}'
                rate_key = f'rate_{item.id}'
                if quantity_key in request.POST and rate_key in request.POST:
                    try:
                        quantity = parse_german_number((request.POST[quantity_key]))
                        unit = request.POST[unit_key]
                        rate = parse_german_number((request.POST[rate_key]))
                        item.quantity = quantity
                        item.unit = unit
                        item.rate = rate
                        item.total = quantity * rate
                        item.save()
                        print(f"Received item details: quantity={quantity}, unit={unit}, rate={rate}")  # Debug line
                    except (ValueError, Item.DoesNotExist):
                        continue

        # Save the contract with the updated VAT and additional fee percentages
        contract.save()

        return JsonResponse({'status': 'success'})

    return JsonResponse({'error': 'Invalid request'}, status=400)


def load_item_users(request):
    item_id = request.GET.get('item_id')
    item = get_object_or_404(Item, id=item_id)
    users = list(item.users.values_list('id', flat=True))
    return JsonResponse({'users': users})

def load_item_budget(request):
    item_id = request.GET.get('item_id')
    item = get_object_or_404(Item, id=item_id)
    budget = item.budget
    return JsonResponse({'budget': budget})

from django.shortcuts import get_object_or_404, redirect
from django.http import JsonResponse
from .models import Item

def add_users_to_item(request, item_id):
    item = get_object_or_404(Item, id=item_id)
    if request.method == 'POST':
        user_ids = request.POST.getlist('users')
        item.users.set(user_ids)
        item.save()  # This will trigger the cascading update
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'fail'}, status=400)

def get_project_users(request):
    project_id = request.GET.get('project_id')
    print(f"Received project_id: {project_id}")  # Debugging line
    try:
        project = Project.objects.get(id=project_id)
        users = project.user.all()
        user_list = [{'id': user.id, 'username': user.username} for user in users]
        return JsonResponse({'project_users': user_list})
    except Project.DoesNotExist:
        return JsonResponse({'error': 'Project not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
@csrf_exempt
@login_required
def delete_contract(request, contract_id):
    if request.method == 'POST':
        contract = get_object_or_404(Contract, id=contract_id)
        project_id = contract.project_set.first().id  # Assuming a contract belongs to at least one project
        contract.delete()
        messages.success(request, 'Contract deleted successfully.')
        return redirect(reverse('edit_project', args=[project_id]))
    return JsonResponse({'error': 'Invalid request'}, status=400)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # Correct namespace handling

def set_bullet(paragraph):
    """
    Apply a bullet point style to a paragraph programmatically.
    """

    pPr = paragraph._element.get_or_add_pPr()  # Get paragraph properties
    numPr = OxmlElement("w:numPr")  # Create numPr element

    ilvl = OxmlElement("w:ilvl")  # List indentation level
    ilvl.set(qn("w:val"), "0")  # Top-level bullet point

    numId = OxmlElement("w:numId")  # Ensure this is used for bullets
    numId.set(qn("w:val"), "1")  # Word uses predefined bullet style with numId = 1

    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)  # Attach the bullet point definition to the paragraph

    # ✅ Ensure correct indentation for multi-line bullets
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")  # Adjust left indentation (lower = closer to margin)
    ind.set(qn("w:hanging"), "360")  # Hanging indent ensures alignment of wrapped text
    pPr.append(ind)

import requests
from decimal import Decimal

def extract_hoai_details(contract):
    """
    Extracts and formats HOAI data from the contract for document generation.
    """
    hoai_data = contract.hoai_data if contract.hoai_data else {}
    is_hoai_contract = hoai_data.get("is_hoai_contract", False)

    # **Extract General HOAI Fields**
    service_profile_name = hoai_data.get("service_profile_name", "Unknown")
    honorarzone = hoai_data.get("honorarzone", "Unknown")
    honorarsatz = hoai_data.get("honorarsatz", "Unknown")
    honorarsatz_factor = hoai_data.get("honorarsatz_factor", "0")

    # **Extract Monetary Values (Keep German Locale Format)**
    baukonstruktionen = hoai_data.get("baukonstruktionen", "0")
    technische_anlagen = hoai_data.get("technischeAnlagen", "0")
    anrechenbare_kosten = hoai_data.get("anrechenbareKosten", "0")

    interpolated_basishonorarsatz = hoai_data.get("interpolated_basishonorarsatz", "0")
    interpolated_oberer_honorarsatz = hoai_data.get("interpolated_oberer_honorarsatz", "0")
    grundhonorar = hoai_data.get("grundhonorar", "0")

    zuschlag_amount = hoai_data.get("zuschlag_amount", "0")
    zuschlag  = hoai_data.get("zuschlag", "0")

    # **Extract LP Breakdown**
    lp_breakdown_actual = hoai_data.get("lp_breakdown_actual", {})
    lp_selected_values = hoai_data.get("lp_values", {})

    # **Extract HOAI Interpolation Details**
    hoai_interpolation = hoai_data.get("interpolation", {})
    lower_bound_cost = hoai_interpolation.get("lower_bound_cost", "0")
    upper_bound_cost = hoai_interpolation.get("upper_bound_cost", "0")
    lower_bound_von = hoai_interpolation.get("lower_bound_von", "0")
    upper_bound_von = hoai_interpolation.get("upper_bound_von", "0")
    lower_bound_bis = hoai_interpolation.get("lower_bound_bis", "0")
    upper_bound_bis = hoai_interpolation.get("upper_bound_bis", "0")

    return {
        "is_hoai_contract": is_hoai_contract,
        "service_profile_name": service_profile_name,
        "honorarzone": honorarzone,
        "honorarsatz": honorarsatz,
        "honorarsatz_factor": honorarsatz_factor,
        "baukonstruktionen": baukonstruktionen,
        "technische_anlagen": technische_anlagen,
        "anrechenbare_kosten": anrechenbare_kosten,
        "interpolated_basishonorarsatz": interpolated_basishonorarsatz,
        "interpolated_oberer_honorarsatz": interpolated_oberer_honorarsatz,
        "grundhonorar": grundhonorar,
        "lp_breakdown_actual": lp_breakdown_actual,
        "lp_values":lp_selected_values,
        "lower_bound_cost": lower_bound_cost,
        "upper_bound_cost": upper_bound_cost,
        "lower_bound_von": lower_bound_von,
        "upper_bound_von": upper_bound_von,
        "lower_bound_bis": lower_bound_bis,
        "upper_bound_bis": upper_bound_bis,
        "zuschlag_amount": zuschlag_amount,
        "zuschlag" : zuschlag,
    }

def insert_html_to_docx(html_content, doc, placeholder):
    """
    Convert HTML content to formatted DOCX content and replace the placeholder text in the Word document.

    Args:
        html_content (str): The HTML string to convert.
        doc (docx.Document): The Word document object.
        placeholder (str): Placeholder text to replace in the Word document.
    """
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html_content, "html.parser")

    # Helper: Apply bullet point style
    def set_bullet(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()

        numPr = OxmlElement('w:numPr')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')

        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')

        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

    # Helper: Apply formatting
    def apply_formatting(run, element):
        run.font.name = "Neue Hans Kendrick"
        run.font.size = Pt(8)
        if element.name == "strong":
            run.bold = True
        if element.name == "u":
            run.underline = True

    # Find the paragraph with placeholder
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            placeholder_paragraph = p
            break

    if not placeholder_paragraph:
        print(f"Placeholder '{placeholder}' not found in the document.")
        return

    # Clear the placeholder text
    placeholder_paragraph.clear()

    # Parse and insert content
    for para in soup.find_all(['p', 'ul', 'ol']):
        if para.name == 'p':
            # Handle <p><br/></p> or empty paragraph → insert one hard break (empty paragraph)
            if not para.get_text(strip=True) and not para.find(['img', 'strong', 'u']):
                placeholder_paragraph.insert_paragraph_before()
                continue

            # Normal <p> with content
            paragraph = placeholder_paragraph.insert_paragraph_before()
            paragraph.paragraph_format.left_indent = Inches(1.75)

            for element in para.contents:
                if element.name in ['strong', 'u']:
                    run = paragraph.add_run(element.get_text())
                    apply_formatting(run, element)
                elif element.name == 'br':
                    run = paragraph.add_run()
                    run.add_break()
                elif element.name is None:
                    run = paragraph.add_run(element.strip())
                    apply_formatting(run, para)

        elif para.name in ['ul', 'ol']:
            for li in para.find_all('li'):
                paragraph = placeholder_paragraph.insert_paragraph_before()
                set_bullet(paragraph)
                paragraph.paragraph_format.left_indent = Inches(2.25)

                for element in li.contents:
                    if element.name in ['strong', 'u']:
                        run = paragraph.add_run(element.get_text())
                        apply_formatting(run, element)
                    elif element.name is None:
                        run = paragraph.add_run(element.strip())
                        apply_formatting(run, li)

                placeholder_paragraph = paragraph  # advance reference


def generate_word_document(request, contract_id):
    print('request:', request)
    template_name = request.GET.get('template_name', 'KOST_De.docx')
    valid_until = request.GET.get('valid_until')
    terms_conditions = request.GET.get('terms_conditions')
    include_scope_of_work = request.GET.get('include_scope_of_work')

    print(f"Template Name: {template_name}")
    print(f"Valid Until: {valid_until}")

    contract = get_object_or_404(Contract, id=contract_id)
    project = contract.project_set.first()
    client = project.client_name

    template_path = os.path.join(r'C:\Users\BCK-CustomApp\Documents\GitHub\Django-HTMX-Finance-App\templates\estimates', template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    doc = DocxTemplate(template_path)

    client_name = getattr(client, 'client_name', 'Unknown')
    firm_name = getattr(client, 'firm_name', 'Unknown')
    street_address = getattr(client, 'street_address', 'Unknown')
    city = getattr(client, 'city', 'Unknown')
    postal_code = getattr(client, 'postal_code', 'Unknown')
    country = getattr(client.country, 'name', 'Unknown') if hasattr(client, 'country') else 'Unknown'

    scope_of_work_html = contract.scope_of_work
    print('scope_of_work_html:', scope_of_work_html)

    # Initialize contract sections and sums
    contract_sections = []
    lp_sections = []
    sum_of_items = Decimal(0)
    sum_of_all_lps = Decimal(0)  # ✅ New variable for LP sum
    section_counter = 1
    is_english_template = template_name in ['BCK_En.docx', 'Kost_En.docx']

    sum_of_items_for_nachlass = Decimal(0)  # Tracks only eligible totals
    nachlass_item_serials = []              # Stores item serials as strings



    hoai_details = extract_hoai_details(contract)
    grundhonorar = parse_german_number(hoai_details["grundhonorar"]) if hoai_details["grundhonorar"] != "0" else Decimal(0)

    # **Process Sections, Separating LP and Non-LP Sections**
    for section in sorted(contract.section.all(), key=lambda s: getattr(s, 'order', 0)):
        section_name = section.section_name
        section_total = Decimal(0)
        items = []
        item_counter = 1

        for item in sorted(section.Item.all(), key=lambda i: getattr(i, 'order', 0)):
            item_total = Decimal(item.quantity) * Decimal(item.rate)

            # Adjust unit names if using an English template
            unit = item.unit
            if is_english_template:
                unit = {'Psch': 'Lumpsum', 'Stk': 'Piece', 'Std': 'Hour'}.get(unit, unit)

            if item.description:
                description = item.description,
            items.append({
                'Item_serial': f"{section_counter}.{item_counter}",
                'Item_name': item.Item_name,
                'description': item.description if item.description else '',
                'quantity': f"{item.quantity:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'unit': unit,
                'rate': f"{Decimal(item.rate):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'total': f"{item_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            })
            section_total += item_total
            if not section.exclude_from_nachlass:
                sum_of_items_for_nachlass += item_total
                nachlass_item_serials.append(f"{section_counter}.{item_counter}")

            item_counter += 1

        # **If Section Name Contains "LP", Store in LP Sections**
        lp_match = re.search(r"LP(\d+)", section_name, re.IGNORECASE)
        if lp_match:
            lp_key = f"lp{lp_match.group(1)}"
            lp_value = hoai_details["lp_values"].get(lp_key, "0")
            actual_lp_value = hoai_details["lp_breakdown_actual"].get(lp_key, "0")



            lp_percentage = Decimal(lp_value) if lp_value != "0" else Decimal(0)
            lp_amount = (lp_percentage / Decimal(100)) * grundhonorar  # ✅ Correct LP calculation

            sum_of_all_lps += lp_amount  # ✅ Add to LP sum

            lp_sections.append({
                'lp_name': section_name,
                'lp_percentage': f"{lp_percentage:.2f}%",
                'lp_amount': f"{lp_amount:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'actual_lp_value': f"{actual_lp_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'Item': items
            })
        else:
            contract_sections.append({
                'section_serial': section_counter,
                'section_name': section_name,
                'net_section': f"{section_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'Item': items
            })
            sum_of_items += section_total
            section_counter += 1

    # **Calculate Contract Totals (Including LP Sections)**
    additional_fee_percentage = Decimal(contract.additional_fee_percentage)

    if contract.hoai_data:
        if sum_of_items  > 0:
            additional_fee_value = ((sum_of_all_lps+sum_of_items) * additional_fee_percentage) / Decimal(100)
        else:
            additional_fee_value = (sum_of_all_lps * additional_fee_percentage) / Decimal(100)
    else:   
        additional_fee_value = (sum_of_items * additional_fee_percentage) / Decimal(100)

    zuschlag_amount = float(parse_german_number(hoai_details["zuschlag_amount"]) if hoai_details["zuschlag_amount"] != "0" else Decimal(0))

    grundhonorar_without_zuschlag = float(grundhonorar) - zuschlag_amount

    nachlass_percentage = Decimal(contract.nachlass_percentage or 0)

    errechnetes_Gesamthonorar  = sum_of_items + sum_of_all_lps + additional_fee_value 
    nachlass_value = (nachlass_percentage / Decimal(100)) * sum_of_items_for_nachlass
    net_contract = sum_of_items + sum_of_all_lps + additional_fee_value - nachlass_value


    vat_percentage = float(contract.vat_percentage) / 100
    tax = float(net_contract) * vat_percentage
    gross_contract = net_contract + Decimal(tax)

    # **Prepare Context for DOCX Template**
    context = {
        'contract_name': contract.contract_name,
        'contract_no': contract.contract_no,
        'project_name': project.project_name,
        'project_no': project.project_no,
        'client_name': client_name,
        'client_firm': firm_name,
        'client_address': f"{street_address}\n{postal_code} {city} \n{country}",

        'contract_sections': contract_sections,
        'sum_of_items': f"{sum_of_items:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
        'sum_of_all_lps': f"{sum_of_all_lps:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  # ✅ LP sum
        
        'net_contract': f"{net_contract:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
        'errechnetes_Gesamthonorar': f"{errechnetes_Gesamthonorar:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),

        'tax': f"{Decimal(tax):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
        'gross_contract': f"{gross_contract:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
        'today_date': date.today().strftime('%d.%m.%Y'),
        'valid_until': valid_until if not valid_until else date.fromisoformat(valid_until).strftime('%d.%m.%Y'),
        
        'vat_percentage': f"{vat_percentage * 100:.2f}",
        'terms_conditions': terms_conditions,
        'include_scope_of_work': include_scope_of_work,
        
        "is_hoai_contract": hoai_details["is_hoai_contract"],
        "service_profile_name": hoai_details["service_profile_name"],
        "honorarzone": hoai_details["honorarzone"],
        "honorarsatz": hoai_details["honorarsatz"],
        "honorarsatz_factor": hoai_details["honorarsatz_factor"],
        "baukonstruktionen": format_german_number(hoai_details["baukonstruktionen"]),
        "technische_anlagen": format_german_number(hoai_details.get("technische_anlagen", "0")),

        "anrechenbare_kosten": format_german_number(hoai_details["anrechenbare_kosten"]),
        "interpolated_basishonorarsatz": hoai_details["interpolated_basishonorarsatz"],
        "interpolated_oberer_honorarsatz": hoai_details["interpolated_oberer_honorarsatz"],
        "grundhonorar": hoai_details["grundhonorar"],
        "lower_bound_cost":hoai_details["lower_bound_cost"] ,
        "upper_bound_cost":hoai_details["upper_bound_cost"] ,
        "lower_bound_von": hoai_details["lower_bound_von"] ,
        "upper_bound_von": hoai_details["upper_bound_von"] ,
        "lower_bound_bis": hoai_details["lower_bound_bis"] ,
        "upper_bound_bis": hoai_details["upper_bound_bis"],
        "zuschlag_amount" : format_german_number(zuschlag_amount),
        "grundhonorar_without_zuschlag":format_german_number(grundhonorar_without_zuschlag),
        "zuschlag_value" : hoai_details["zuschlag"],
    }

    if additional_fee_percentage > 0:
        context.update({
            'additional_fee_percentage': f"{additional_fee_percentage:.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'additional_fee_value': f"{additional_fee_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        })

    if nachlass_value != 0:
        context.update({
            'nachlass_value' : f"{nachlass_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'nachlass_percentage' : f"{nachlass_percentage:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'nachlass_applied_items': ", ".join(nachlass_item_serials),

        })
         
    if lp_sections:
        context['lp_sections'] = lp_sections

    import pprint
    pprint.pprint(context)

    doc.render(context)

    if include_scope_of_work == 'on':
        insert_html_to_docx(scope_of_work_html, doc, placeholder="[[SCOPE_OF_WORK]]")

    company_name = template_name.split("_")[0]
    project_short_name = project.project_name.split()[0]  # Gets the first word
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file_name = f"{contract.contract_no} {company_name} {project_short_name} AN {contract.contract_name}.docx"
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'

    doc.save(response)
    return response


import locale

def format_german_number(number):
    """
    Converts a number into German locale format.
    Example: 1000000 → "1.000.000,00"
    """
    number = Decimal(number)
    locale.setlocale(locale.LC_NUMERIC, 'de_DE.UTF-8')  # Set German locale
    return locale.format_string("%.2f", number, grouping=True).replace(',', 'X').replace('.', ',').replace('X', '.')

from django.contrib import messages

@login_required
def create_invoice(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    contracts = project.contract.all()

    selected_contract = contracts.first() if contracts else None
    additional_fee_percentage = selected_contract.additional_fee_percentage if selected_contract else 0
    vat_percentage = selected_contract.vat_percentage if selected_contract else 0

    if request.method == 'POST':
        form = InvoiceForm(request.POST, project=project)
        if form.is_valid():
            invoice = form.save(commit=False)
            invoice.project = project
            invoice.provided_quantities = json.loads(request.POST.get('provided_quantities'))

            # Set is_cumulative field based on form input
            is_cumulative = request.POST.get('is_cumulative', 'false') == 'true'
            invoice.is_cumulative = is_cumulative

            invoice.save()

            messages.success(request, 'Invoice created successfully.')
            return HttpResponseRedirect(reverse('edit_project', args=[project_id]) + '?tab=invoices')
        else:
            return JsonResponse({'status': 'error', 'errors': form.errors})
    else:
        form = InvoiceForm(project=project)

    return render(request, 'tracker/create_invoice.html', {
        'form': form,
        'project': project,
        'contracts': contracts,
        'additional_fee_percentage': additional_fee_percentage,
        'vat_percentage': vat_percentage
    })

def calculate_available_quantity(project, item):
    all_invoices = Invoice.objects.filter(project=project)
    total_quantity_invoiced = sum(
        json.loads(invoice.provided_quantities).get(str(item.id), {}).get('quantity', 0)
        for invoice in all_invoices
    )
    return item.quantity - total_quantity_invoiced

# View for deleting an invoice
@login_required
def delete_invoice(request, invoice_id):
    invoice = get_object_or_404(Invoice, id=invoice_id)
    project_id = invoice.project.id  # Assuming an invoice belongs to a project
    invoice.delete()

    # Redirect to the edit project page with the invoices tab open
    return redirect(reverse('edit_project', args=[project_id]) + '?tab=invoices')

@login_required
def view_invoice(request, invoice_id):
    try:
        invoice = get_object_or_404(Invoice, id=invoice_id)
        project = invoice.project
        contract = invoice.contract

        # Parse provided quantities from the invoice
        try:
            provided_quantities_data = invoice.provided_quantities
        except ValueError as e:
            print(f"Error parsing provided_quantities: {e}")
            return JsonResponse({'error': 'Invalid data format for provided quantities'}, status=400)

        provided_quantities = []
        sum_of_items = Decimal('0.00')
        nachlass_applicable_sum = Decimal('0.00')


        for item_id, details in provided_quantities_data.items():
            try:
                item = Item.objects.get(id=item_id)
                section = Section.objects.filter(Item=item, contract=contract).first()
                section_name = section.section_name if section else "Unknown Section"

                rate = Decimal(details['rate'])
                quantity = Decimal(details['quantity'])
                total = rate * quantity
                if section and not section.exclude_from_nachlass:
                    nachlass_applicable_sum += total

                sum_of_items += total

                provided_quantities.append({
                    'section_name': section_name,
                    'item_name': item.Item_name,
                    'unit': item.unit,
                    'rate': str(rate),
                    'quantity': str(quantity),
                    'total': str(total),
                    'exclude_from_nachlass': section.exclude_from_nachlass if section else False,
                })
            except Item.DoesNotExist:
                print(f"Item with id {item_id} does not exist.")
                rate = Decimal(details['rate'])
                quantity = Decimal(details['quantity'])
                total = rate * quantity
                sum_of_items += total

                provided_quantities.append({
                    'section_name': 'Unknown Section',
                    'item_name': f'Unknown Item (ID: {item_id})',
                    'unit': 'Unknown Unit',
                    'rate': str(rate),
                    'quantity': str(quantity),
                    'total': str(total),
                })

        # Base financial calculations
        additional_fee_percentage = Decimal(contract.additional_fee_percentage or 0)
        additional_fee_value = (sum_of_items * additional_fee_percentage) / Decimal(100)
        vat_percentage = Decimal(contract.vat_percentage or 0)
        nachlass_percentage = Decimal(contract.nachlass_percentage or 0)

        additional_fee_value = (sum_of_items * additional_fee_percentage) / Decimal(100)
        nachlass_value = (nachlass_applicable_sum * nachlass_percentage) / Decimal(100)
        invoice_net = sum_of_items + additional_fee_value - nachlass_value
        tax_value = (invoice_net * vat_percentage) / Decimal(100)
        invoice_gross = invoice_net + tax_value

        data = {
            'project_name': project.project_name,
            'contract_name': contract.contract_name,
            'additional_fee_percentage': str(additional_fee_percentage),
            'additional_fee_value': str(additional_fee_value),
            'provided_quantities': provided_quantities,
            'invoice_net': str(invoice_net),
            'tax_value': str(tax_value),
            'invoice_gross': str(invoice_gross),
            'vat_percentage': str(vat_percentage),
            'amount_received': invoice.amount_received,
            'date_of_payment': invoice.date_of_payment.isoformat() if invoice.date_of_payment else None,
            'nachlass_percentage': str(nachlass_percentage),
            'nachlass_value': str(nachlass_value),
        }

        #  Adjust for cumulative invoices
        if invoice.invoice_type in ['AR', 'SR', 'ZR']:
            previous_invoices = Invoice.objects.filter(
                project=project,
                contract=contract,
                created_at__lt=invoice.created_at,
                invoice_type__in=['AR', 'SR', 'ZR']
            )

            total_previous_net = Decimal('0.00')
            total_previous_tax = Decimal('0.00')

            for prev in previous_invoices:
                prev_net = Decimal(prev.invoice_net)
                prev_tax = (prev_net * vat_percentage) / Decimal(100)
                total_previous_net += prev_net
                total_previous_tax += prev_tax

            current_invoice_net = invoice_net - total_previous_net
            current_invoice_tax = tax_value - total_previous_tax
            current_invoice_gross = current_invoice_net + current_invoice_tax

            data.update({
                'current_invoice_net': str(current_invoice_net),
                'current_invoice_tax': str(current_invoice_tax),
                'current_invoice_gross': str(current_invoice_gross),
            })

        return JsonResponse(data)

    except Http404 as e:
        print(f"Invoice or related data not found: {e}")
        return JsonResponse({'error': 'Invoice not found'}, status=404)

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return JsonResponse({'error': 'An unexpected error occurred'}, status=500)
        
from django.utils.dateparse import parse_date
import pprint
from django.http import FileResponse
import socket

def download_invoice(request, invoice_id):
    # Fetch the invoice, project, and contract
    invoice = get_object_or_404(Invoice, id=invoice_id)
    project = invoice.project
    contract = invoice.contract
    client = project.client_name

    # Get the template name and date range from the request
    template_name = request.GET.get('invoice_template_name', 'inv_BCK_De.docx')
    print(f"Template name from modal: {template_name}")

    from_date_str = request.GET.get('from_date')
    to_date_str = request.GET.get('to_date')

    from_date = parse_date(from_date_str) if from_date_str else None
    to_date = parse_date(to_date_str) if to_date_str else None

    # Check if template is English
    is_english_template = template_name in ['inv_BCK_En.docx', 'inv_Kost_En.docx']

    # Extract HOAI details
    hoai_details = extract_hoai_details(contract)
    grundhonorar = parse_german_number(hoai_details["grundhonorar"]) if hoai_details["grundhonorar"] != "0" else Decimal(0)
    zuschlag_amount = float(parse_german_number(hoai_details["zuschlag_amount"]) if hoai_details["zuschlag_amount"] != "0" else Decimal(0))
    grundhonorar_without_zuschlag = float(grundhonorar) - zuschlag_amount

    # Build section order mapping from contract
    section_order = {
        section.section_name: idx + 1
        for idx, section in enumerate(contract.section.order_by('order'))
    }

    # Initialize section storage
    contract_sections_dict = {}
    lp_sections = []
    sum_of_items = Decimal('0.00')
    sum_of_all_lps = Decimal('0.00')

    # Initialize invoice data
    nachlass_applicable_sum = Decimal('0.00')
    nachlass_item_serials = []

    # Build item order mapping per section
    section_item_order = {
        section.id: {
            item.id: idx + 1
            for idx, item in enumerate(section.Item.order_by('order'))
        }
        for section in contract.section.all()
    }

    # Process provided quantities
    provided_quantities = invoice.provided_quantities
    for item_id, details in provided_quantities.items():
        item = get_object_or_404(Item, id=item_id)
        section = item.section_set.first()
        section_name = section.section_name if section else "Unknown Section"
        exclude_from_nachlass = section.exclude_from_nachlass if section else False  # 💡 Key line
        section_serial = section_order.get(section_name, 0)
        item_total = Decimal(details['quantity']) * Decimal(details['rate'])

        unit = item.unit
        if is_english_template:
            unit = {'Psch': 'Lumpsum', 'Stk': 'Piece', 'Std': 'Hour'}.get(unit, unit)

        lp_match = re.search(r"LP(\d+)", section_name, re.IGNORECASE)
        if lp_match:
            lp_key = f"lp{lp_match.group(1)}"
            lp_value = hoai_details["lp_values"].get(lp_key, "0")
            actual_lp_value = hoai_details["lp_breakdown_actual"].get(lp_key, "0")
            lp_percentage = Decimal(lp_value) if lp_value != "0" else Decimal(0)
            lp_amount = (Decimal(details['quantity']) / Decimal(100)) * grundhonorar
            sum_of_all_lps += lp_amount

            item_index = section_item_order.get(section.id, {}).get(item.id, 0)
            item_serial = f"{section_serial}.{item_index}"

            # Apply Nachlass check for LP items (if ever needed)
            if not exclude_from_nachlass:
                nachlass_applicable_sum += lp_amount
                nachlass_item_serials.append(item_serial)

            lp_sections.append({
                'lp_name': section_name,
                'lp_percentage': f"{lp_percentage:.2f}%",
                'lp_amount': f"{lp_amount:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'actual_lp_value': f"{actual_lp_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'Item': [{
                    'Item_name': item.Item_name,
                    'Item_serial': item_serial,
                    'unit': unit,
                    'rate': f"{Decimal(details['rate']):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                    'quantity': f"{details['quantity']:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                    'total': f"{item_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                }]
            })
        else:
            if section_name not in contract_sections_dict:
                contract_sections_dict[section_name] = {
                    'section_serial': section_serial,
                    'section_name': section_name,
                    'net_section': Decimal('0.00'),
                    'Item': []
                }

            section_data = contract_sections_dict[section_name]
            item_index = section_item_order.get(section.id, {}).get(item.id, 0)
            item_serial = f"{section_serial}.{item_index}"
            # item_serial = f"{section_data['section_serial']}.{len(section_data['Item']) + 1}"

            section_data['Item'].append({
                'Item_name': item.Item_name,
                'Item_serial': item_serial,
                'unit': unit,
                'rate': f"{Decimal(details['rate']):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'quantity': f"{Decimal(details['quantity']):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
                'total': f"{item_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            })

            section_data['net_section'] += item_total
            sum_of_items += item_total

            # Track for Nachlass only if not excluded
            if not exclude_from_nachlass:
                nachlass_applicable_sum += item_total
                nachlass_item_serials.append(item_serial)

    # Sort and finalize contract sections
    contract_sections = []
    for section_name in sorted(contract_sections_dict.keys(), key=lambda n: section_order.get(n, 999)):
        section = contract_sections_dict[section_name]
        section['net_section'] = f"{section['net_section']:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        contract_sections.append(section)

    # Calculate totals
    additional_fee_percentage = Decimal(contract.additional_fee_percentage)


    if contract.hoai_data:
        if sum_of_items  > 0:
            additional_fee_value = ((sum_of_all_lps+sum_of_items) * additional_fee_percentage) / Decimal(100)
        else:
            additional_fee_value = (sum_of_all_lps * additional_fee_percentage) / Decimal(100)

    else:   
        additional_fee_value = (sum_of_items * additional_fee_percentage) / Decimal(100)

    nachlass_percentage = Decimal(contract.nachlass_percentage or 0)

    nachlass_value = (nachlass_applicable_sum * nachlass_percentage) / Decimal(100)

    invoice_net = sum_of_items + sum_of_all_lps + additional_fee_value 
    vat_percentage = Decimal(contract.vat_percentage) / Decimal(100)
    vat_percentage_display = Decimal(contract.vat_percentage)
    tax_value = invoice_net * vat_percentage
    invoice_gross = invoice_net + tax_value

    # Calculate cumulative invoice adjustments
    previous_invoices = Invoice.objects.filter(
        project=project,
        contract=contract,
        created_at__lt=invoice.created_at
    ).order_by('created_at')

    total_invoice_gross = Decimal('0.00')
    total_invoice_net = Decimal('0.00')
    total_invoice_tax = Decimal('0.00')
    total_amount_paid = Decimal('0.00')

    previous_invoices_data = []
    for inv in previous_invoices:
        if inv.invoice_type not in ['AR', 'SR', 'ZR']:
            continue
        inv_net = Decimal(inv.invoice_net)
        inv_tax = inv_net * vat_percentage
        inv_gross = inv_net + inv_tax
        inv_paid = Decimal(inv.amount_received)

        total_invoice_gross += inv_gross
        total_invoice_net += inv_net
        total_amount_paid += inv_paid
        total_invoice_tax += inv_tax

        previous_invoices_data.append({
            'invoice_title': inv.title,
            'invoice_type': inv.invoice_type,
            'created_at': timezone.localtime(inv.created_at).strftime('%d.%m.%Y'),
            'invoice_net': f"{inv_net:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'invoice_tax%': vat_percentage,
            'invoice_tax': f"{inv_tax:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'invoice_gross': f"{inv_gross:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'amount_paid': f"{inv_paid:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
        })

    current_invoice_net = invoice_net - total_invoice_net - nachlass_value
    current_invoice_tax = current_invoice_net * vat_percentage
    current_invoice_gross = current_invoice_net + current_invoice_tax
    invoice_tobepaid = total_invoice_gross - total_amount_paid
    previous_ar_count = previous_invoices.filter(invoice_type='AR').count()


    # Prepare context for template rendering
    context = {
        'contract_name': contract.contract_name,
        'contract_no' : contract.contract_no,
        'project_name': project.project_name,
        'project_no': project.project_no,
        'client_name': client.client_name if client else "Unknown",
        'client_firm': client.firm_name if client else "Unknown",
        'client_address': f"{client.street_address}\n{client.postal_code} {client.city} \n{client.country.name}",
        
        'created_at': timezone.localtime(invoice.created_at).strftime('%d.%m.%Y'),  # German format with time
        
        'invoice_title': invoice.title,

        "additional_fee_percentage":additional_fee_percentage,
        "additional_fee_value" :  f"{additional_fee_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),

        'contract_sections': contract_sections,  # Organized by section
        'sum_of_items': f"{sum_of_items:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  
        'lp_sections': lp_sections,  # Organized by LP
        'sum_of_all_lps': f"{sum_of_all_lps:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        
        'invoice_net': f"{invoice_net:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  
        'tax': f"{tax_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        'invoice_gross': f"{invoice_gross:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        
        'current_invoice_net' :  f"{current_invoice_net:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  
        'current_invoice_gross' :  f"{current_invoice_gross:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  
        'current_invoice_tax' :  f"{current_invoice_tax:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),  

        'vat_percentage': vat_percentage,  
        'vat_percentage_display': vat_percentage_display, 

        "is_hoai_contract": hoai_details["is_hoai_contract"],
        "service_profile_name": hoai_details["service_profile_name"],
        "honorarzone": hoai_details["honorarzone"],
        "honorarsatz": hoai_details["honorarsatz"],
        "honorarsatz_factor": hoai_details["honorarsatz_factor"],
        "baukonstruktionen": format_german_number(hoai_details["baukonstruktionen"]),
        "technische_anlagen": format_german_number(hoai_details["technische_anlagen"]),
        "anrechenbare_kosten": format_german_number(hoai_details["anrechenbare_kosten"]),
        "interpolated_basishonorarsatz": hoai_details["interpolated_basishonorarsatz"],
        "interpolated_oberer_honorarsatz": hoai_details["interpolated_oberer_honorarsatz"],
        "grundhonorar": hoai_details["grundhonorar"],
        "lower_bound_cost":hoai_details["lower_bound_cost"] ,
        "upper_bound_cost":hoai_details["upper_bound_cost"] ,
        "lower_bound_von": hoai_details["lower_bound_von"] ,
        "upper_bound_von": hoai_details["upper_bound_von"] ,
        "lower_bound_bis": hoai_details["lower_bound_bis"] ,
        "upper_bound_bis": hoai_details["upper_bound_bis"],
        "zuschlag_amount" : format_german_number(zuschlag_amount),
        "grundhonorar_without_zuschlag":format_german_number(grundhonorar_without_zuschlag),
        "zuschlag_value" : hoai_details["zuschlag"],

        'previous_invoices': previous_invoices_data, 
        'total_invoice_gross': f"{total_invoice_gross:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        'total_invoice_net': f"{total_invoice_net:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        'total_invoice_tax': f"{total_invoice_tax:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        
        'total_amount_paid': f"{total_amount_paid:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        'invoice_tobepaid': f"{invoice_tobepaid:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'), 
        
        'invoice_type': invoice.invoice_type,
        'from_date': from_date.strftime('%d.%m.%Y') if from_date else None,
        'to_date': to_date.strftime('%d.%m.%Y') if to_date else None,
        'client_firm': client.firm_name
    }

    # Sort nachlass_item_serials numerically
    sorted_nachlass_items = sorted(
        nachlass_item_serials,
        key=lambda x: tuple(map(int, x.split('.')))
    )
    if nachlass_value != 0:
        context.update({
            'nachlass_value': f"{nachlass_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'nachlass_percentage': f"{nachlass_percentage:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'),
            'nachlass_applied_items': ", ".join(sorted_nachlass_items)
        })


    # Load and render template
    template_path = os.path.join(r'C:\Users\BCK-CustomApp\Documents\GitHub\Django-HTMX-Finance-App\templates\invoices', template_name)
    doc = DocxTemplate(template_path)
    doc.render(context)
    print(context)
    # Build HTTP response with correct file name
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    project_short_name = project.project_name.split()[0]
    company_identifier = "BCK" if "BCK" in template_name else "KOST"
    new_filename = f"{invoice.title} {company_identifier} {project_short_name} {invoice.invoice_type} {contract.contract_name}.docx"
    response['Content-Disposition'] = f'attachment; filename={new_filename}'
    doc.save(response)

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(
    buffer,
    as_attachment=True,
    filename=new_filename,
    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
)

def record_payment(request, invoice_id):
    invoice = get_object_or_404(Invoice, id=invoice_id)

    if request.method == 'POST':
        try:
            amount_received = float(request.POST.get('amount_received', 0))
            invoice.amount_received = amount_received

            date_str = request.POST.get('date_of_payment')
            if date_str:
                try:
                    invoice.date_of_payment = datetime.strptime(date_str, '%Y-%m-%d').date()
                except ValueError:
                    messages.warning(request, 'Invalid date format. Date of payment was not saved.')

            invoice.save()
            messages.success(request, 'Payment recorded successfully.')
            return redirect(reverse('edit_project', args=[invoice.project.id]) + '?tab=invoices')
        except (ValueError, TypeError):
            return JsonResponse({'error': 'Invalid input for amount or date.'}, status=400)

    return JsonResponse({'error': 'Invalid request'}, status=400)


from django.http import JsonResponse
from .models import EstimateSettings, Project

def get_new_contract_number(request, project_id):
    try:
        # Fetch the current consecutive number from EstimateSettings
        settings = EstimateSettings.objects.first()
        if not settings:
            return JsonResponse({"error": "Estimate settings not configured."}, status=400)

        consecutive_no = settings.consecutive_start_no

        # Fetch the project
        project = Project.objects.get(id=project_id)
        project_no = project.project_no

        # Combine consecutive_no and project_no to create a unique contract number
        contract_no = f"{consecutive_no}-{project_no}"

        return JsonResponse({"contract_no": contract_no})

    except Project.DoesNotExist:
        return JsonResponse({"error": "Project not found."}, status=404)

from bs4 import BeautifulSoup

@csrf_exempt
def update_scope(request, contract_id):
    if request.method == 'POST':
        scope_of_work = request.POST.get('scope_of_work', '')

        # Validate HTML
        try:
            soup = BeautifulSoup(scope_of_work, 'html.parser')
            sanitized_html = str(soup)  # Ensure valid HTML structure
            print(f"Sanitized HTML: {sanitized_html}")
        except Exception as e:
            print(f"HTML parsing error: {e}")
            return JsonResponse({'status': 'error', 'message': 'Invalid HTML content'}, status=400)

        try:
            contract = Contract.objects.get(id=contract_id)
            contract.scope_of_work = sanitized_html
            contract.save()
            return JsonResponse({'status': 'success', 'message': 'Scope updated successfully'})
        except Exception as e:
            print(f"Database save error: {e}")
            return JsonResponse({'status': 'error', 'message': f'Error saving contract: {str(e)}'}, status=500)

from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
import pandas as pd
from .models import ServiceProfile
from .serializers import ServiceProfileSerializer

class ServiceProfileUploadView(APIView):
    """Allows users to upload an Excel file"""
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        serializer = ServiceProfileSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "File uploaded successfully!", "data": serializer.data}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ServiceProfileListView(APIView):
    """Lists all uploaded service profiles, including LP breakdowns"""
    def get(self, request, *args, **kwargs):
        profiles = ServiceProfile.objects.all()
        
        response_data = [
            {
                "id": profile.id,
                "name": profile.name,
                "no_of_Honarzone": profile.no_of_Honarzone,  
                "lp_breakdown": profile.lp_breakdown  # Add LP breakdown
            }
            for profile in profiles
        ]

        return Response(response_data, status=200)

class HOAICalculationView(APIView):
    """Performs calculations based on the selected Excel file and chargeable costs"""
    
    def post(self, request, *args, **kwargs):
        profile_id = request.data.get('service_profile_id')
        cost_input = float(request.data.get('chargeable_costs'))
        fee_zone = request.data.get('fee_zone') 
        profile = get_object_or_404(ServiceProfile, id=profile_id)
        excel_path = profile.excel_file.path

        # Read Excel file
        df = pd.read_excel(excel_path)

        # Find the lower and upper bound rows
        lower_bound = df[df["Anrechenbare Kosten (€)"] <= cost_input].iloc[-1]
        upper_bound = df[df["Anrechenbare Kosten (€)"] > cost_input].iloc[0]

        # Extract values
        a = lower_bound["Anrechenbare Kosten (€)"]
        aa = upper_bound["Anrechenbare Kosten (€)"]
        
        b = lower_bound[f"Honorarzone {fee_zone} (von)"]
        bb = upper_bound[f"Honorarzone {fee_zone} (von)"]
        
        c = lower_bound[f"Honorarzone {fee_zone} (bis)"]
        cc = upper_bound[f"Honorarzone {fee_zone} (bis)"]

        # Perform linear interpolation
        honor_from = b + ((cost_input - a) * (bb - b) / (aa - a))
        honor_to = c + ((cost_input - a) * (cc - c) / (aa - a))

        response_data = {
            "service_profile": profile.name,
            "chargeable_costs": cost_input,
            "fee_zone": fee_zone,
            "interpolation": {
                "lower_bound_cost": a,
                "upper_bound_cost": aa,
                "lower_bound_von": b,
                "upper_bound_von": bb,
                "lower_bound_bis": c,
                "upper_bound_bis": cc,
            },
            "calculated_fee": {
                "honor_from": round(honor_from, 2),
                "honor_to": round(honor_to, 2),
            }
        }

        return Response(response_data, status=status.HTTP_200_OK)

import os
import uuid

def generate_revit_shared_parameter_file(project):
    """Generate a Revit shared parameter file with project-specific details."""

    file_content = f"""# This is a Revit shared parameter file.
# Do not edit manually.
*META    VERSION    MINVERSION
META    2    1
*GROUP    ID    NAME
GROUP    1    Project Information
*PARAM    GUID    NAME    DATATYPE    DATACATEGORY    GROUP    VISIBLE    DESCRIPTION    USERMODIFIABLE    HIDEWHENNOVALUE
PARAM    {uuid.uuid4()}    BCK_Project Status    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Projektkürzel    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_BH_Straße    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Client Name    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Project Issue Date    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_IN_Mengen Kommentare    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Organizations Name    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Project Name    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Project Number    TEXT        1    1        1    0
PARAM    {uuid.uuid4()}    BCK_Organization Description    TEXT        1    1        1    0
"""

    # Replace placeholders with actual project details
    file_content = file_content.replace("BCK_Project Status", project.status)
    file_content = file_content.replace("BCK_Projektkürzel", project.project_no)
    file_content = file_content.replace("BCK_BH_Straße", project.project_address)
    file_content = file_content.replace("BCK_Client Name", project.client_name.client_name if project.client_name else "Unknown Client")
    file_content = file_content.replace("BCK_Project Issue Date", str(project.client_name.client_mail if project.client_name else "Unknown Date"))
    file_content = file_content.replace("BCK_IN_Mengen Kommentare", "Default Comment")
    file_content = file_content.replace("BCK_Organizations Name", "BCK Architecture")
    file_content = file_content.replace("BCK_Project Name", project.project_name)
    file_content = file_content.replace("BCK_Project Number", project.project_no)
    file_content = file_content.replace("BCK_Organization Description", "Architectural Firm")

    # Define file path
    file_path = f"media/projects/{project.project_no}_shared_params.txt"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Write to file
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(file_content)

    return file_path

@login_required
def get_first_invoice_mode(request):
    contract_id = request.GET.get('contract_id')
    if contract_id:
        first_invoice = Invoice.objects.filter(contract_id=contract_id).order_by('created_at').first()
        if first_invoice:
            return JsonResponse({
                'first_invoice_exists': True,
                'is_cumulative': first_invoice.is_cumulative
            })
    return JsonResponse({'first_invoice_exists': False})

# views.py - Add view to save project hourly rates
from django.views.decorators.http import require_POST
from django.shortcuts import redirect

@require_POST
def update_project_settings(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    hourly_fields = [
        'executive_management_rate', 'specialist_planner_rate', 'project_management_rate',
        'construction_supervision_rate', 'computational_architect_rate', 'architect_rate',
        'construction_technician_rate', 'draftsman_rate'
    ]
    updated_rates = {}
    for field in hourly_fields:
        value = request.POST.get(field)
        if value:
            try:
                updated_rates[field] = float(value)
            except ValueError:
                continue

    project.hourly_rates_override = updated_rates
    project.save()
    return redirect('edit_project', project_id=project.id)

from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages

@login_required
def reset_project_hourly_rates(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    print('reset project hours')
    print(project)
    if project.hourly_rates_override:
        project.hourly_rates_override = {}
        project.save()
        messages.success(request, 'Die benutzerdefinierten Stundensätze wurden zurückgesetzt.')

    return redirect('edit_project', project_id=project.id)